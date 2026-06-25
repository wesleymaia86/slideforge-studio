"""DOCX text extraction using python-docx."""
import io

from docx import Document

from .base import ParseResult


def parse_docx(data: bytes, preview_rows: int = 20) -> ParseResult:
    doc = Document(io.BytesIO(data))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    tables_count = len(doc.tables)

    schema = [
        {"name": "paragraph_text", "dtype": "string", "nullable": False, "sampleValues": paragraphs[:5]},
        {"name": "tables_count", "dtype": "integer", "nullable": False, "sampleValues": [tables_count]},
    ]

    preview = [
        {"index": i, "paragraph": text}
        for i, text in enumerate(paragraphs[:preview_rows])
    ]

    return ParseResult(
        row_count=len(paragraphs),
        column_count=1,
        sheet_names=["Document"],
        schema=schema,
        preview=preview,
    )
